from io import BytesIO
from pathlib import Path

from fastapi import UploadFile


MAX_FILE_SIZE = 15 * 1024 * 1024
SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


class ExtractError(Exception):
    """Raised when uploaded document text cannot be extracted."""


async def extract_text_from_upload(file: UploadFile) -> str:
    if not file.filename:
        raise ExtractError("Vui lòng chọn một file tài liệu.")

    extension = Path(file.filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ExtractError("Định dạng chưa được hỗ trợ. Hãy dùng file .txt, .pdf hoặc .docx.")

    content = await file.read()
    if not content:
        raise ExtractError("File đang trống, không có nội dung để tóm tắt.")
    if len(content) > MAX_FILE_SIZE:
        raise ExtractError("File quá lớn. Vui lòng dùng file nhỏ hơn 15 MB.")

    if extension == ".txt":
        text = _extract_txt(content)
    elif extension == ".pdf":
        text = _extract_pdf(content)
    else:
        text = _extract_docx(content)

    normalized = _normalize_text(text)
    if not normalized:
        raise ExtractError("Không trích xuất được văn bản từ file này.")
    return normalized


def _extract_txt(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1258", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ExtractError("Không đọc được mã hóa của file .txt.")


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ExtractError("Thiếu thư viện pypdf. Hãy cài dependencies trong requirements.txt.") from exc

    try:
        reader = PdfReader(BytesIO(content))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:  # pypdf raises several parser-specific exceptions.
        raise ExtractError("Không đọc được nội dung file PDF.") from exc


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ExtractError("Thiếu thư viện python-docx. Hãy cài dependencies trong requirements.txt.") from exc

    try:
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        return "\n".join(paragraphs + table_cells)
    except Exception as exc:
        raise ExtractError("Không đọc được nội dung file DOCX.") from exc


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").split("\n")]
    return "\n".join(line for line in lines if line).strip()
